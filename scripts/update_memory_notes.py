from copy import deepcopy

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_BREAK, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm
from docx.text.paragraph import Paragraph


DOC_PATH = "/Users/lauramachado/Desktop/Trabajo/cyberguide-tfg/docs/Memoria_CyberGuide_Borrador.docx"


def paragraph_text(paragraph):
    return "".join(run.text for run in paragraph.runs)


def clear_paragraph(paragraph):
    p = paragraph._element
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)


def set_paragraph_text(paragraph, text):
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    return paragraph


def delete_paragraph(paragraph):
    element = paragraph._element
    parent = element.getparent()
    parent.remove(element)


def insert_paragraph_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        run = new_para.add_run(text)
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    return new_para


def insert_page_break_after(paragraph):
    page_break = insert_paragraph_after(paragraph, style="Normal")
    run = page_break.add_run()
    run.add_break(WD_BREAK.PAGE)
    return page_break


def find_paragraph(doc, text, style_name=None, occurrence=1):
    count = 0
    for paragraph in doc.paragraphs:
        if paragraph_text(paragraph).strip() != text:
            continue
        if style_name and paragraph.style.name != style_name:
            continue
        count += 1
        if count == occurrence:
            return paragraph
    raise ValueError(f"Paragraph not found: {text!r} style={style_name!r} occurrence={occurrence}")


def add_index_line(after_paragraph, text, indent_cm):
    paragraph = insert_paragraph_after(after_paragraph, text, style="Normal")
    paragraph.paragraph_format.left_indent = Cm(indent_cm)
    paragraph.paragraph_format.first_line_indent = Cm(0)
    return paragraph


def add_highlighted_paragraph(after_paragraph, text, style="Normal"):
    return insert_paragraph_after(after_paragraph, text, style=style)


def ensure_section_break_before(doc, target_text):
    target = find_paragraph(doc, target_text, "Heading 1")
    prev = target._p.getprevious()
    if prev is None:
        return
    prev_para = Paragraph(prev, target._parent)
    ppr = prev_para._p.get_or_add_pPr()
    if ppr.find(qn("w:sectPr")) is not None:
        return
    body_sectpr = doc._element.body.sectPr
    new_sectpr = deepcopy(body_sectpr)
    type_el = new_sectpr.find(qn("w:type"))
    if type_el is None:
        type_el = OxmlElement("w:type")
        new_sectpr.insert(0, type_el)
    type_el.set(qn("w:val"), "nextPage")
    ppr.append(new_sectpr)


def rebuild_general_index(doc):
    heading = find_paragraph(doc, "Índice general", "Heading 1")
    figures_heading = find_paragraph(doc, "Índice de figuras", "Heading 1")

    current = heading._p.getnext()
    while current is not None and current != figures_heading._p:
        next_node = current.getnext()
        current.getparent().remove(current)
        current = next_node

    entries = [
        ("1. Introducción", 0.0),
        ("1.1. Motivación y planteamiento del proyecto", 0.8),
        ("1.2. Contexto del problema", 0.8),
        ("1.3. Estructura de la memoria", 0.8),
        ("2. Objetivos", 0.0),
        ("2.1. Objetivo general", 0.8),
        ("2.2. Objetivos específicos", 0.8),
        ("2.3. Alcance del proyecto", 0.8),
        ("2.4. Elementos fuera de alcance", 0.8),
        ("3. Tecnologías utilizadas", 0.0),
        ("3.1. Conceptos previos y glosario técnico", 0.8),
        ("3.2. Tecnologías principales", 0.8),
        ("3.2.1. Backend y orquestación del sistema", 1.6),
        ("3.2.2. Modelo generativo y ejecución local", 1.6),
        ("3.2.3. Recuperación semántica y almacenamiento vectorial", 1.6),
        ("3.2.4. Procesamiento documental y soporte multimodal", 1.6),
        ("3.2.5. Interfaz web y experiencia de uso", 1.6),
        ("3.3. Corpus documental, criterios de selección y trazabilidad de fuentes", 0.8),
        ("3.4. Tecnologías auxiliares y despliegue", 0.8),
        ("3.5. Estructura del proyecto y organización de archivos", 0.8),
        ("4. Diseño e implementación del proyecto", 0.0),
        ("4.1. Metodología seguida y planificación temporal", 0.8),
        ("4.2. Sprint 1: definición del proyecto y preparación del entorno", 0.8),
        ("4.3. Sprint 2: diseño técnico y primer backend funcional", 0.8),
        ("4.4. Sprint 3: primera validación end-to-end y mejora del núcleo RAG", 0.8),
        ("4.5. Sprint 4: ampliación funcional con interfaz, PDF y OCR", 0.8),
        ("4.6. Sprint 5: seguridad, evaluación automática e iteración del backend", 0.8),
        ("4.7. Sprint 6: integración frontend-backend y evolución del producto", 0.8),
        ("5. Experimentos y validación", 0.0),
        ("5.1. Metodología de evaluación", 0.8),
        ("5.2. Primera validación end-to-end", 0.8),
        ("5.3. Validación multimodal y comportamiento seguro", 0.8),
        ("5.4. Resultados del benchmark automático", 0.8),
        ("5.5. Discusión de resultados y limitaciones", 0.8),
        ("6. Conclusiones", 0.0),
        ("6.1. Consecución de objetivos", 0.8),
        ("6.2. Aplicación de lo aprendido", 0.8),
        ("6.3. Lecciones aprendidas", 0.8),
        ("6.4. Trabajos futuros", 0.8),
        ("7. Bibliografía", 0.0),
        ("8. Anexos", 0.0),
    ]

    cursor = heading
    for text, indent in entries:
        cursor = add_index_line(cursor, text, indent)


def add_blank_page_between_summary_and_index(doc):
    summary_break = find_paragraph(doc, "Summary", "Heading 1")
    next_nodes = []
    current = summary_break._p.getnext()
    while current is not None:
        para = Paragraph(current, summary_break._parent)
        if paragraph_text(para).strip() == "Índice general":
            break
        next_nodes.append(para)
        current = current.getnext()
    if len(next_nodes) < 3:
        tail = next_nodes[-1] if next_nodes else summary_break
        insert_page_break_after(tail)


def update_intro_structure(doc):
    intro = find_paragraph(doc, "1. Introducción", "Heading 1")
    next_para = Paragraph(intro._p.getnext(), intro._parent)
    if paragraph_text(next_para).strip() != "1.1. Motivación y planteamiento del proyecto":
        insert_paragraph_after(intro, "1.1. Motivación y planteamiento del proyecto", style="Heading 2")

    set_paragraph_text(find_paragraph(doc, "1.1. Contexto del problema", "Heading 2"), "1.2. Contexto del problema")
    set_paragraph_text(find_paragraph(doc, "1.2. Estructura de la memoria", "Heading 2"), "1.3. Estructura de la memoria")

    structure_para = find_paragraph(
        doc,
        "La memoria se organiza en varios capítulos que recorren el proyecto desde la contextualización del problema hasta la evaluación final. Antes de ellos se incluye un apartado preliminar de definiciones y conceptos clave pensado para facilitar la lectura del resto del documento.",
        "Normal",
    )
    set_paragraph_text(
        structure_para,
        "Tras la portada, la memoria incorpora un bloque preliminar formado por agradecimientos, resumen, summary e índices. A partir de ahí, el contenido principal se organiza en varios capítulos que recorren el proyecto desde la motivación inicial y la definición de objetivos hasta el desarrollo técnico, la evaluación y las conclusiones finales.",
    )


def expand_glossary(doc):
    spa_anchor = find_paragraph(
        doc,
        "React: biblioteca de JavaScript utilizada para construir la interfaz del frontend mediante componentes reutilizables y facilitar una interfaz conversacional modular.",
        "List Bullet",
    )
    paragraph = insert_paragraph_after(
        spa_anchor,
        "SPA (Single Page Application): tipo de aplicación web en la que la navegación principal se resuelve en el cliente sin recargar por completo la página, algo especialmente útil para una experiencia de chat fluida.",
        style="List Bullet",
    )

    api_anchor = find_paragraph(
        doc,
        "API REST: interfaz basada en peticiones HTTP que permite comunicar frontend y backend de manera estructurada.",
        "List Bullet",
    )
    endpoint = insert_paragraph_after(
        api_anchor,
        "Endpoint: ruta concreta expuesta por una API para realizar una acción determinada, como consultar el estado del servicio o enviar una pregunta al asistente.",
        style="List Bullet",
    )
    prompt = insert_paragraph_after(
        endpoint,
        "Prompt: instrucción o plantilla textual que guía al modelo de lenguaje para responder con el tono, el formato y las restricciones deseadas.",
        style="List Bullet",
    )
    insert_paragraph_after(
        prompt,
        "Grounding: grado en que la respuesta generada se mantiene fiel a la evidencia recuperada y evita introducir afirmaciones no respaldadas por las fuentes.",
        style="List Bullet",
    )


def add_corpus_section(doc):
    old_aux = find_paragraph(doc, "3.3. Tecnologías auxiliares y despliegue", "Heading 2")
    new_section = insert_paragraph_after(
        Paragraph(old_aux._p.getprevious(), old_aux._parent),
        "3.3. Corpus documental, criterios de selección y trazabilidad de fuentes",
        style="Heading 2",
    )

    p1 = insert_paragraph_after(
        new_section,
        "La especificidad de CyberGuide no depende solo del modelo generativo o de la arquitectura RAG, sino del corpus documental que alimenta el sistema. Frente a un asistente generalista, cuyo conocimiento es amplio pero difuso y difícil de auditar, este proyecto se apoya en un conjunto acotado de fuentes públicas, institucionales y trazables. Esa decisión condiciona la calidad de la respuesta, pero también constituye una de las principales justificaciones del backend: el valor del sistema reside en cómo selecciona, organiza, recupera y presenta evidencia útil para el dominio concreto de la ciberseguridad en pymes y trabajo autónomo.",
        style="Normal",
    )
    c1 = insert_paragraph_after(
        p1,
        "La selección del corpus siguió varios criterios: origen oficial o altamente fiable, acceso público, utilidad práctica para el público objetivo, lenguaje suficientemente claro para poder ser reinterpretado por el asistente y relevancia respecto a riesgos frecuentes como contraseñas, copias de seguridad, teletrabajo, respuesta ante incidentes o gestión de crisis. Este filtrado previo fue necesario para evitar que el sistema se limitara a responder con información genérica o poco alineada con las necesidades reales del usuario final.",
        style="Normal",
    )
    c2 = insert_paragraph_after(
        c1,
        "La primera base documental operativa del prototipo se construyó con seis documentos oficiales de INCIBE, que durante la validación inicial generaron 225 fragmentos indexados en la colección local `cyberguide`. A partir de esa base se ensayó el comportamiento del sistema en consultas sobre copias de seguridad, teletrabajo seguro, respuesta ante incidentes y gestión de contraseñas. El anexo 8.1 recoge el inventario inicial del corpus y la documentación del proyecto conserva, además, un registro más amplio de fuentes candidatas y de futuras ampliaciones.",
        style="Normal",
    )
    insert_paragraph_after(
        c2,
        "Este planteamiento aporta dos ventajas relevantes para el TFG. En primer lugar, permite explicar con claridad de dónde salen los datos que usa el asistente y por qué se consideran adecuados para el dominio elegido. En segundo lugar, hace que el sistema sea escalable y auditable: ampliar CyberGuide no consiste en pedir al modelo que “sepa más”, sino en incorporar nuevas fuentes, reingestar el corpus y comprobar de forma reproducible cómo cambia la recuperación y la respuesta final.",
        style="Normal",
    )

    set_paragraph_text(old_aux, "3.4. Tecnologías auxiliares y despliegue")
    set_paragraph_text(
        find_paragraph(doc, "3.4. Estructura del proyecto y organización de archivos", "Heading 2"),
        "3.5. Estructura del proyecto y organización de archivos",
    )


def fix_objective_numbering(doc):
    numbered = [
        "Seleccionar y organizar un corpus documental compuesto por fuentes públicas, oficiales y trazables relacionadas con ciberseguridad para pequeñas organizaciones y trabajo autónomo.",
        "Implementar una arquitectura RAG local capaz de realizar ingesta documental, fragmentación del contenido, generación de embeddings, indexación vectorial y recuperación semántica.",
        "Desarrollar una interfaz conversacional que permita consultar el sistema en lenguaje natural y acceder a respuestas útiles sin exigir conocimientos técnicos avanzados al usuario.",
        "Incorporar soporte multimodal controlado para el análisis puntual de documentos PDF e imágenes o capturas, siempre dentro de una respuesta textual y prudente.",
        "Diseñar mecanismos de seguridad orientados a reducir comportamientos arriesgados en casos como phishing, suplantación o solicitudes de credenciales.",
        "Evaluar el comportamiento del prototipo mediante casos de prueba representativos, atendiendo a la relevancia de la recuperación, la calidad del grounding, la utilidad práctica y la gestión de la incertidumbre.",
    ]
    for idx, text in enumerate(numbered, start=1):
        para = find_paragraph(doc, text, "List Number")
        para.style = "Normal"
        para.paragraph_format.left_indent = Cm(0.7)
        para.paragraph_format.first_line_indent = Cm(-0.4)
        set_paragraph_text(para, f"{idx}. {text}")


def enrich_methodology(doc):
    anchor = find_paragraph(
        doc,
        "Este planteamiento permitió mantener controlado el alcance del proyecto y documentar con claridad la evolución del backend, la interfaz, el corpus, la capa OCR y la evaluación. Por ello, en esta memoria el trabajo se presenta mediante sprints semanales, ya que esa forma refleja mejor cómo se organizó realmente el progreso de CyberGuide.",
        "Normal",
    )
    insert_paragraph_after(
        anchor,
        "Aunque no existió un equipo Scrum formal ni se aplicaron todas sus ceremonias de manera estricta, sí hubo artefactos de seguimiento equivalentes a pequeña escala: un registro diario de avances, un log de decisiones, notas de validación y una priorización continua del siguiente objetivo técnico. Esa combinación permitió sostener un desarrollo individual pero disciplinado, muy cercano en la práctica a una dinámica ágil de ingeniería.",
        style="Normal",
    )


def main():
    doc = Document(DOC_PATH)

    rebuild_general_index(doc)
    add_blank_page_between_summary_and_index(doc)
    update_intro_structure(doc)
    expand_glossary(doc)
    add_corpus_section(doc)
    fix_objective_numbering(doc)
    enrich_methodology(doc)
    ensure_section_break_before(doc, "1. Introducción")

    doc.save(DOC_PATH)


if __name__ == "__main__":
    main()
